from flask import Flask, render_template, request, redirect, url_for, session
import pandas as pd
import os
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor
import matplotlib.pyplot as plt
import io
import base64
from docx import Document
from pptx import Presentation
import pdfkit

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # For session management

@app.route('/')
def index():
    return render_template('index.html')

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files['file']
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    if file.filename.endswith('.csv'):
        df = pd.read_csv(file_path)
    elif file.filename.endswith('.xlsx'):
        df = pd.read_excel(file_path)

    # Save the dataframe in session for future use
    session['data'] = df.to_dict()

    return render_template('select_model.html', columns=df.columns)


@app.route('/train_model', methods=['POST'])
def train_model():
    df = pd.DataFrame(session['data'])  # Retrieve dataset from session
    model_type = request.form['model']
    column = request.form['column']

    X = df.drop(column, axis=1)  # Features
    y = df[column]  # Target

    if model_type == 'linear_regression':
        model = LinearRegression()
    elif model_type == 'decision_tree':
        model = DecisionTreeRegressor()

    model.fit(X, y)

    # Save model and data for use in graph generation
    session['model'] = model
    session['X'] = X.to_dict()
    session['y'] = y.to_dict()
    session['column'] = column
    session['model_type'] = model_type

    return redirect(url_for('select_graph'))

@app.route('/generate_graph', methods=['POST'])
def generate_graph():
    X = pd.DataFrame(session['X'])
    y = pd.Series(session['y'])
    graph_type = request.form['graph_type']

    plt.figure()
    if graph_type == 'scatter':
        plt.scatter(X, y)
    elif graph_type == 'line':
        plt.plot(X, y)

    # Save graph image as base64
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    graph_url = base64.b64encode(img.getvalue()).decode()

    # Store graph data in session for possible deletion or report generation
    session['graph_url'] = graph_url

    return render_template('display_graph.html', graph_url=graph_url)

@app.route('/add_graph', methods=['POST'])
def add_graph():
    return redirect(url_for('select_model'))  # Redirect back to model selection for new graph

@app.route('/delete_graph', methods=['POST'])
def delete_graph():
    session.pop('graph_url', None)  # Remove current graph from session
    return redirect(url_for('select_model'))

@app.route('/write_summary', methods=['GET'])
def write_summary():
    return render_template('write_summary.html')

@app.route('/save_summary', methods=['POST'])
def save_summary():
    session['summary'] = request.form['summary']
    return redirect(url_for('generate_report'))

@app.route('/generate_report', methods=['GET'])
def generate_report():
    summary = session.get('summary', '')
    graph_url = session.get('graph_url', '')

    # Generate the report in the chosen format (let's use DOCX here as an example)
    doc = Document()
    doc.add_heading('Analysis Report', 0)
    doc.add_paragraph(summary)

    # Add the graph image to the report
    if graph_url:
        img_data = base64.b64decode(graph_url)
        with open('graph.png', 'wb') as f:
            f.write(img_data)

        doc.add_picture('graph.png')

    doc.save('report.docx')
    return send_file('report.docx', as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
